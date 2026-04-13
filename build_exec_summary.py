from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EXECUTIVE SUMMARY')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Value to Victory, LLC')
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Confidential \u2014 April 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# Info box table
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_left = table.cell(0, 0)
cell_right = table.cell(0, 1)

left_lines = [
    ("Management:", True),
    ("CEO: Shawn E. Decker", False),
    ("VP Fitness/Coaching: Cameron Decker", False),
    ("Operations: Kyle Hegstrum", False),
    ("Client Success: Sandra Aldridge", False),
    ("", False),
    ("Industry: AI-Powered Personal Development / CoachTech", True),
    ("Number of Employees: 4 (core team)", False),
    ("Amount of Financing Sought: $1,750,000 (Series Seed)", False),
    ("Current Investors: Self-funded (Founder capital)", False),
    ("Use of Funds: Engineering (31%), Marketing (29%), Enterprise (17%), Ops (14%), Legal/IP (9%)", False),
]

right_lines = [
    ("Contact:", True),
    ("valuetovictory@gmail.com", False),
    ("valuetovictory.com", False),
    ("", False),
    ("Entity: Virginia LLC", False),
    ("Founded: 2024", False),
    ("Platform Launch: Live (2025)", False),
    ("Monthly Infrastructure: $8\u201310K/mo", False),
    ("Current MRR: $29 (early stage)", False),
]

cell_left.text = ""
for i, (text, bold) in enumerate(left_lines):
    if i == 0:
        p = cell_left.paragraphs[0]
    else:
        p = cell_left.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(1)

cell_right.text = ""
for i, (text, bold) in enumerate(right_lines):
    if i == 0:
        p = cell_right.paragraphs[0]
    else:
        p = cell_right.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(1)

doc.add_paragraph()

# Section helper
def add_section(title, content):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)
    for para_text in content.split('\n\n'):
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            run.font.size = Pt(10)
    doc.add_paragraph()


add_section('Company Description',
"Value to Victory, LLC is an AI-powered personal development platform built on the proprietary "
"P.I.N.K. Framework (People, Influence, Numbers, Knowledge, with Time as a 0.1x\u20132.0x multiplier). "
"The company delivers assessment-driven coaching, relationship tools, and faith-based growth systems "
"through a fully integrated technology stack: 88 API endpoints, 39 database tables, and 14 addressable "
"market verticals.\n\n"
"VTV is the only platform that combines multi-dimensional life assessment, AI-generated coaching, "
"couples/dating compatibility, fitness integration (via FitCarna merger), and a Book-to-System conversion "
"pipeline \u2014 all under one roof. Founded by U.S. Navy veteran and licensed real estate appraiser "
"Shawn E. Decker, VTV transforms personal development from generic advice into precision-guided, "
"data-driven transformation.")

add_section('Problem / Opportunity',
"The personal development industry is a $243B+ market fragmented across disconnected point solutions. "
"BetterUp charges $500+/session for human coaching. Gallup offers static assessments with no action plan. "
"Match Group reduces relationships to photo swiping. Peloton addresses only physical fitness.\n\n"
"No platform connects life assessment, coaching, relationships, fitness, faith, and financial growth "
"into a single system with measurable outcomes. Consumers spend thousands across 4\u20136 separate platforms "
"and still lack a unified roadmap. VTV solves this with one assessment, one platform, and one coaching "
"engine addressing all five pillars simultaneously.")

add_section('Our Solution',
"VTV\u2019s P.I.N.K. Assessment scores users across 5 life pillars (80\u2013120 data points), then generates "
"personalized coaching recommendations through a proprietary Cross-Pillar Impact Matrix (20 directional "
"relationships between pillars). The Time Multiplier (0.1x\u20132.0x) collapses all scores when time "
"management fails, creating urgency and driving coaching conversions.\n\n"
"Autonomous AI coaching delivers morning motivation and evening accountability emails at zero marginal "
"cost. The platform includes couples mode (partner invitations, shared dashboards, compatibility scoring), "
"dating/matching (standards-based, not photo swiping), and a Book-to-System pipeline that converts any "
"self-help book into a live coaching system.\n\n"
"FitCarna merger adds 300+ existing fitness clients, 12K+ lbs lost, and closed-loop physical pillar coaching "
"from Day 1. Immediate revenue: $12K+/mo from existing client base, scaling to $65K/mo via gym affiliate "
"coaching partnerships.")

add_section('Market',
"Total Addressable Market: $243B+ across 14 verticals.\n\n"
"Primary markets (live now): Personal coaching ($15.2B), Relationship/dating ($9.6B), Corporate wellness "
"($61B), Real estate tech ($20B), Faith-based development ($5B+).\n\n"
"Expansion verticals (by 2028): Enterprise HR assessments, coach certification/licensing, Book-to-System "
"publishing, white-label church programs, AI tutoring, and real estate appraisal API.\n\n"
"Market growth rate: CoachTech sector growing 15\u201320% annually. AI coaching subsector growing 35%+ as "
"enterprises seek scalable alternatives to human-only coaching.")

add_section('Clients',
"1. Growth-Minded Professionals (ages 28\u201355): Earn $60K\u2013$150K, seeking work-life balance, "
"priced out of premium coaching. Entry: free assessment, convert at $29/mo.\n\n"
"2. Couples & Families: Partners seeking compatibility insights and shared growth dashboards. "
"Entry: couples assessment, convert at $47/mo.\n\n"
"3. Faith-Based Communities: Churches and small groups wanting Scripture-integrated personal development. "
"Entry: free devotionals, convert to membership.\n\n"
"4. Enterprise HR/L&D Teams: Companies seeking scalable employee development with privacy controls. "
"Entry: team pilot, convert at $8\u2013$15/seat/mo.\n\n"
"5. Authors & Coaches: Content creators wanting to convert their book into a live coaching system. "
"Entry: Book-to-System license at $2,500\u2013$10,000 + recurring.")

add_section('Revenue Model',
"12 compounding revenue streams across 3 categories:\n\n"
"Subscriptions (B2C):\n"
"\u2022 VictoryPath: $29/mo ($290/yr) \u2014 Core assessment + AI coaching\n"
"\u2022 Value Builder: $47/mo ($470/yr) \u2014 Couples mode + dating + group features\n"
"\u2022 Victory VIP: $497/mo ($4,970/yr) \u2014 1:1 coaching + priority support + all access\n\n"
"Platform & Licensing (B2B):\n"
"\u2022 Book-to-System License: $2,500\u2013$10,000 setup + $200\u2013500/mo + 10\u201315% revenue share\n"
"\u2022 Enterprise Seats: $8\u2013$15/seat/mo\n"
"\u2022 Coach Certification: $2,500\u2013$5,000 per cohort\n"
"\u2022 White-Label / Church Programs: $5\u2013$500/mo\n"
"\u2022 RE Appraisal API: $0.10\u2013$0.50/query\n\n"
"FitCarna Coaching:\n"
"\u2022 1:1 Coaching (Carna Method): $200\u2013$500/mo\n"
"\u2022 8-Week Challenge: $100/participant\n"
"\u2022 PowerHour Classes: $50\u2013$100/mo\n"
"\u2022 Personal Coaching: $125\u2013$500/hr")

add_section('Marketing & Sales Channels',
"Digital-First Distribution:\n"
"\u2022 Free assessment funnel (valuetovictory.com) \u2014 primary acquisition channel\n"
"\u2022 Social media: Facebook, Instagram, YouTube, LinkedIn, X (all active)\n"
"\u2022 SEO/content marketing via Obsidian-powered blog and devotional content\n"
"\u2022 Automated email sequences: nurture, morning coaching, evening accountability\n"
"\u2022 n8n automation: 7 active workflows for lead capture, CRM sync, onboarding\n"
"\u2022 HubSpot CRM: real-time contact sync with 6 custom properties\n\n"
"Direct Sales:\n"
"\u2022 1:1 coaching consultations (Calendly booking)\n"
"\u2022 Enterprise pilot programs (direct outreach)\n"
"\u2022 Book-to-System licensing (author/coach partnerships)\n"
"\u2022 Church/faith community partnerships\n\n"
"Strategic:\n"
"\u2022 FitCarna cross-referral (assessment identifies fitness gaps, refers to coaching)\n"
"\u2022 Startup Warrior founder cohort (investor pipeline)\n"
"\u2022 Conference speaking and workshop delivery")

add_section('Competitors in Every Vertical \u2014 VTV Connects Them',
"6+ vertical competitors exist \u2014 none connects assessment, coaching, relationships, fitness, and faith "
"into one platform. VTV is the connective tissue:\n\n"
"\u2022 Wiley/DiSC ($1.97B): 4-dimensional workplace assessments. VTV: 5 pillars + AI + couples.\n"
"\u2022 BetterUp ($4.7B): Human coaching at $500+/session. VTV: AI coaching at $29/mo.\n"
"\u2022 Match Group ($8.3B): Photo-based swiping. VTV: whole-life compatibility scoring.\n"
"\u2022 Gallup ($500M+): 34 static themes, no action plan. VTV: dynamic scoring + Time Multiplier.\n"
"\u2022 Peloton ($1.97B): Fitness only. VTV: FitCarna + 4 non-physical pillars.\n"
"\u2022 Noom ($4.0B): Weight loss only. VTV: 5 pillars, not just physical.\n\n"
"Key differentiator: VTV\u2019s Time Multiplier (provisional patent filed) and Cross-Pillar Impact Matrix "
"have zero prior art. The true moat is compounding user data across all five pillars \u2014 every assessment "
"trains the model, making the platform more valuable over time.")

add_section('Our Team',
"Shawn E. Decker \u2014 Founder & CEO\n"
"U.S. Navy Veteran (Aviation Maintenance, 1998\u20132002). Licensed Real Estate Appraiser, 23+ years. "
"Published author (2 books). Created the P.I.N.K. Framework and built the entire platform architecture.\n\n"
"Cameron Decker \u2014 VP Fitness & Coaching (FitCarna)\n"
"Leads the Carna Method coaching system. 300+ existing clients, 12K+ lbs lost. $12K+/mo Day 1 revenue, "
"scaling to $65K/mo via gym affiliate coaching partnerships. Manages fitness pillar integration.\n\n"
"Kyle Hegstrum \u2014 Operations\n"
"Operational support and business development. Client onboarding and partnership coordination.\n\n"
"Sandra Aldridge \u2014 Client Success\n"
"Early adopter and client success lead. Scored 171 on P.I.N.K. assessment, gained 10+ hours/week "
"within 90 days. Living proof of the system\u2019s effectiveness.")

add_section('Third-Party Validation',
"External Proof Points:\n"
"\u2022 4\u20135 external companies launching with VTV assessment integration\n"
"\u2022 Sandra Aldridge: Scored 171 on P.I.N.K., gained 10+ hrs/week in 90 days (living case study)\n"
"\u2022 FitCarna: 300+ clients, 12K+ lbs lost, $12K+/mo live revenue from Day 1\n"
"\u2022 Pre-Investment Platform Value: $300\u2013500K (sweat equity equivalent \u2014 88 API endpoints, 39 DB tables, "
"full automation stack built before raising)\n\n"
"Bottom-Up Unit Economics:\n"
"\u2022 Customer Acquisition Cost (CAC): $35\u2013$60\n"
"\u2022 Average Revenue Per User (ARPU): $42/mo (B2C), $208/seat/yr (enterprise)\n"
"\u2022 Retention Target: 85%+\n"
"\u2022 LTV:CAC Ratio: 6:1+\n"
"\u2022 Infrastructure: $8\u201310K/mo at scale")

add_section('Our Needs',
"Raising $1,750,000 in Series Seed funding for 25% equity.\n"
"Pre-money valuation: $5.25M | Post-money: $7.0M\n\n"
"Use of Funds:\n"
"\u2022 Engineering & Product: $540,000 (31%) \u2014 Senior hire, mobile app, AI training, Book-to-System\n"
"\u2022 Sales & Marketing: $500,000 (29%) \u2014 Paid acquisition, content, sales team, partnerships\n"
"\u2022 Enterprise & B2B: $300,000 (17%) \u2014 Enterprise sales, pilot programs, coach certification\n"
"\u2022 Operations & Infra: $250,000 (14%) \u2014 DevOps, security audit, SOC 2, scaling\n"
"\u2022 Reserve & Legal/IP: $160,000 (9%) \u2014 10 patent filings, legal counsel, operating reserve\n\n"
"Investor Protections: 8% preferred return ($140K/yr), 25% quarterly net distributions, "
"1x liquidation preference, weighted-average anti-dilution, board seat, pro-rata rights.")

add_section('What We Want to Achieve',
"12-Month Goals:\n"
"\u2022 Reach $120K MRR ($1.44M ARR)\n"
"\u2022 Launch mobile app (iOS + Android)\n"
"\u2022 Onboard 5 enterprise pilot clients\n"
"\u2022 File 5 provisional patent applications (Tier 1 innovations)\n"
"\u2022 Complete Book-to-System pipeline for first 3 external authors\n"
"\u2022 Scale FitCarna to 500+ active coaching clients\n\n"
"24-Month Goals:\n"
"\u2022 Reach $720K MRR ($8.6M ARR)\n"
"\u2022 Close Series A at $30\u2013$50M valuation\n"
"\u2022 Launch dating/matching vertical\n"
"\u2022 10+ enterprise clients with 1,000+ seats\n"
"\u2022 Coach certification program with 50+ certified coaches\n\n"
"Key Metrics: MRR/ARR (Stripe), assessment conversion rate, NPS (target 70+), enterprise utilization, patent milestones.")

# Financial Projections Table
p = doc.add_paragraph()
run = p.add_run('Financial Projections')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['', 'YEAR 1', 'YEAR 3', 'YEAR 5']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

data = [
    ['SALES (ARR)', '$1,443,600', '$8,634,000', '$28,000,000'],
    ['EBIT', '($180,000)', '$2,800,000', '$11,200,000'],
    ['NET PROFIT / LOSS', '($250,000)', '$2,100,000', '$8,400,000'],
]

for r, row_data in enumerate(data):
    row = table.rows[r + 1]
    for c, val in enumerate(row_data):
        cell = row.cells[c]
        cell.text = val
        for p in cell.paragraphs:
            if c > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    'Note: Year 1 reflects heavy investment in product and team. Breakeven projected Month 14\u201316. '
    'Year 3 assumes enterprise traction and licensing revenue. Year 5 assumes 14-vertical expansion '
    'and potential Series B/IPO positioning.'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

doc.add_paragraph()

# Live Links section
p = doc.add_paragraph()
run = p.add_run('Live Platform Links')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x8B, 0x69, 0x14)

links = [
    ('Main Website', 'https://valuetovictory.com'),
    ('Assessment Platform', 'https://assessment.valuetovictory.com'),
    ('Member Portal', 'https://assessment.valuetovictory.com/member'),
    ('Free Assessment', 'https://assessment.valuetovictory.com/assessment'),
    ('Founder Website', 'https://shawnedecker.com'),
    ('Investor Prospectus (HTML)', 'https://assessment.valuetovictory.com/investor-pitch.html'),
    ('NDA Agreement (HTML)', 'https://assessment.valuetovictory.com/nda-agreement.html'),
    ('Startup Warrior Pitch (HTML)', 'https://assessment.valuetovictory.com/startup-warrior-pitch.html'),
]

for label, url in links:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(url)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Confidential \u2014 Do not distribute without written permission from Value to Victory, LLC')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('valuetovictory@gmail.com')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

out_path = r'C:\Users\Administrator\OneDrive\Desktop\Zillion\vtv-assessment\VTV-Executive-Summary.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
